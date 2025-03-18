from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from translator_agent import TranslatorAgent, llm

class WordTranslator:
    def __init__(self, input_file, output_file, language="french"):
        self.input_file = input_file
        self.output_file = output_file
        self.language = language  # Langue pour la traduction, par défaut "french"
        self.translator_agent = TranslatorAgent(llm, self.language)  # Initialise automatiquement l'agent de traduction avec la langue donnée
        self.doc = Document(input_file)

    @staticmethod
    def get_run_style(run):
        """Retourne un tuple représentant le style d'un run (pour déterminer s'il peut être fusionné)."""
        return (
            run.bold,
            run.italic,
            run.underline,
            run.font.name,
            run.font.size,
            run.font.color.rgb
        )

    @staticmethod
    def extract_hyperlinks_with_positions(paragraph):
        """Extrait les liens hypertextes d'un paragraphe et retourne une liste de (texte, URL, position)."""
        links = []
        for index, child in enumerate(paragraph._element):
            if child.tag == qn("w:hyperlink"):
                r_id = child.get(qn("r:id"))
                if r_id:
                    part = paragraph.part
                    rel = part.rels.get(r_id)
                    if rel:
                        url = rel.target_ref
                        link_text = "".join(
                            node.text for node in child.findall(".//" + qn("w:t")) if node.text
                        )
                        links.append((link_text, url, index))
        return links

    @staticmethod
    def create_hyperlink(paragraph, text, url, position):
        """Ajoute un lien hypertexte cliquable dans un paragraphe."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        r_pr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)

        run.append(r_pr)

        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)

        hyperlink.append(run)
        paragraph._element.insert(position, hyperlink)

    def translate_paragraph(self, para):
        """Fusionne les runs avec le même style, traduit, et les réinsère sans perdre la mise en forme."""
        if not para.runs:
            return

        new_runs = []
        current_text = ""
        current_style = self.get_run_style(para.runs[0])
        links = self.extract_hyperlinks_with_positions(para)

        for run in para.runs:
            if self.get_run_style(run) == current_style:
                current_text += run.text
            else:
                new_runs.append((current_text, current_style))
                current_text = run.text
                current_style = self.get_run_style(run)

        if current_text:
            new_runs.append((current_text, current_style))

        # Traduit le texte en utilisant la méthode `translate` de l'agent de traduction
        translated_runs = [(self.translator_agent.translate(text), style) for text, style in new_runs]

        para.clear()

        for text, style in translated_runs:
            new_run = para.add_run(text)
            new_run.bold, new_run.italic, new_run.underline, new_run.font.name, new_run.font.size, new_run.font.color.rgb = style

        for link_text, url, position in links:
            self.create_hyperlink(para, link_text, url, position)

    def translate_document(self):
        """Traduit l'ensemble du document."""
        for para in self.doc.paragraphs:
            self.translate_paragraph(para)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self.translate_paragraph(para)

        self.doc.save(self.output_file)


def translation_pipeline(input_file="file.docx", output_file="document_traduit2.docx", language="french"):
    """Pipeline pour gérer la traduction du document avec une langue personnalisable."""
    word_translator = WordTranslator(input_file, output_file, language)
    word_translator.translate_document()


if __name__ == "__main__":
  # Exemple d'utilisation (Pas besoin d'initialiser la langue et l'agent de traduction séparément maintenant)
  input_file = "mission_letter_1.docx"
  output_file = f"{input_file.split('.')[0]}_translated.docx"
  translation_pipeline(input_file=input_file, output_file=output_file)

