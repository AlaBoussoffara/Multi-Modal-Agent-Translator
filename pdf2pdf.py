import os
import fitz  # PyMuPDF for PDF processing
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from pdf2docx import Converter  # Convert PDF to DOCX
from docx2pdf import convert  # Convert DOCX back to PDF
from typing import Optional
from translator_agent import TranslatorAgent # Import your existing translator agent

class PDFTranslator:
    """
    Converts a PDF to a structured format, extracts text for translation, 
    and replaces it with the translated text while preserving layout.
    """

    def __init__(self, translator: TranslatorAgent, target_language="french"):
        self.translator = translator
        self.target_language = target_language

    def pdf_to_docx(self, pdf_path: str, docx_path: str):
        """ Converts a PDF file to a DOCX file while preserving layout. """
        print("Converting PDF to DOCX...")
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        print("✅ PDF converted to DOCX")

    def extract_text_from_docx(self, docx_path: str) -> str:
        """ Extracts text from a DOCX file while preserving structure. """
        print("Extracting text from DOCX...")
        doc = Document(docx_path)
        extracted_text = ""

        for para in doc.paragraphs:
            extracted_text += para.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        extracted_text += para.text + "\n"

        return extracted_text.strip()

    def replace_text_in_docx(self, docx_path: str, translated_text: str):
        """ Replaces the extracted text with translated text while preserving styles and handling missing translations. """
        print("Replacing text with translated version while keeping styles...")
        doc = Document(docx_path)
        translated_lines = translated_text.split("\n")
        index = 0

        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                if index < len(translated_lines):  # Ensure index is valid
                    for run in para.runs:
                        translated_run_text = self.translator.translate(run.text)
                        run.text = translated_run_text  # Preserve run styles
                    index += 1

        # Handle tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():  # Skip empty cells
                            if index < len(translated_lines):  # Avoid index error
                                for run in para.runs:
                                    translated_run_text = self.translator.translate(run.text)
                                    run.text = translated_run_text
                                index += 1

        doc.save(docx_path)
        print("✅ DOCX updated with translated text while keeping styles")

    def docx_to_pdf(self, docx_path: str, pdf_path: str):
        """ Converts a modified DOCX back to PDF while keeping formatting. """
        print("Converting DOCX back to PDF...")
        convert(docx_path, pdf_path)
        print(f"✅ Translated PDF saved at: {pdf_path}")

    def translate_pdf(self, pdf_path: str, output_pdf: Optional[str] = None):
        """ Main function to extract, translate, and replace text in a PDF. """
        if not output_pdf:
            output_pdf = pdf_path.replace(".pdf", "_translated(test).pdf")

        docx_path = pdf_path.replace(".pdf", ".docx")

        # Step 1: Convert PDF to DOCX
        self.pdf_to_docx(pdf_path, docx_path)

        # Step 2: Extract text
        extracted_text = self.extract_text_from_docx(docx_path)

        # Step 3: Translate text
        translated_text = self.translator.translate(extracted_text)

        # Step 4: Replace original text with translation
        self.replace_text_in_docx(docx_path, translated_text)

        # Step 5: Convert back to PDF
        self.docx_to_pdf(docx_path, output_pdf)

        # Optional: Clean up intermediate DOCX file
        os.remove(docx_path)

        return output_pdf


# Example Usage
if __name__ == "__main__":
    # Initialize TranslatorAgent (already implemented)
    translator_agent = TranslatorAgent(target_language="english")

    # Initialize PDF translator
    pdf_translator = PDFTranslator(translator=translator_agent, target_language="english")

    # Translate PDF 
    pdf_translator.translate_pdf("Rapport d'audit technique Vaudrimesnil + commentaire EDPR.pdf") 
